"""DOCX parser (python-docx): paragraphs as prose, tables preserved."""

import io

from docx import Document as DocxDocument

from src.knowledge.ingest_errors import ParserError
from src.knowledge.parsing.base import ContentBlock, DocumentFormat, ParsedDocument


class DocxParser:
    def parse(self, data: bytes, *, max_pages: int) -> ParsedDocument:
        """Time: O(paragraphs + table cells). DOCX is reflowable (page_count=1)."""
        try:
            document = DocxDocument(io.BytesIO(data))
        except Exception as exc:
            raise ParserError("file is not a readable DOCX") from exc

        blocks: list[ContentBlock] = []
        for index, paragraph in enumerate(document.paragraphs):
            if paragraph.text.strip():
                blocks.append(ContentBlock(text=paragraph.text, ref={"section": index}))

        table_count = 0
        for table_index, table in enumerate(document.tables):
            rows = [" | ".join(cell.text for cell in row.cells) for row in table.rows]
            if any(row.strip() for row in rows):
                blocks.append(
                    ContentBlock(text="\n".join(rows), ref={"table": table_index}, is_table=True)
                )
                table_count += 1

        ratio = table_count / len(blocks) if blocks else 0.0
        return ParsedDocument(
            blocks=tuple(blocks), page_count=1, format=DocumentFormat.DOCX, table_ratio=ratio
        )
